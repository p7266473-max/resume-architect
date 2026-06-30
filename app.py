"""
Resume Architect Factory
========================
A production-ready Streamlit application that transforms raw career history
into executive-quality resumes using Google Gemini AI and python-docx.
"""

import streamlit as st
import os
import io
import re
import json
import time
from typing import Any, Optional

from google import genai
from google.genai import types

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# CONSTANTS
# ============================================================

APP_TITLE: str = "Resume Architect Factory"
APP_SUBTITLE: str = "Transform raw career history into an executive-quality resume."
APP_ICON: str = "💼"
GEMINI_MODEL: str = "gemini-2.5-flash"
MAX_RETRIES: int = 3
RETRY_DELAY_SECONDS: float = 2.0

THEMES: list[str] = ["Modern-Tech", "Classic-Executive", "ATS-Friendly"]

THEME_COLORS: dict[str, dict[str, RGBColor]] = {
    "Modern-Tech": {
        "heading": RGBColor(0x6C, 0x5C, 0xE7),   # Purple
        "subheading": RGBColor(0x00, 0xB8, 0x94),  # Teal
        "accent": RGBColor(0x0D, 0x95, 0xE8),      # Blue
    },
    "Classic-Executive": {
        "heading": RGBColor(0x1A, 0x1A, 0x2E),     # Navy
        "subheading": RGBColor(0x2D, 0x3A, 0x4A),   # Charcoal
        "accent": RGBColor(0x8B, 0x6F, 0x47),       # Gold-brown
    },
    "ATS-Friendly": {
        "heading": RGBColor(0x00, 0x00, 0x00),     # Black
        "subheading": RGBColor(0x33, 0x33, 0x33),   # Dark gray
        "accent": RGBColor(0x44, 0x44, 0x44),       # Medium gray
    },
}

PLACEHOLDER_TEXT: str = (
    "Civil Engineer with 8 years experience...\n"
    "Managed infrastructure projects...\n"
    "Designed structural systems...\n"
    "Reduced project costs...\n"
    "Bachelor of Engineering..."
)

PASS1_SYSTEM_PROMPT: str = """You are an expert resume writer and career strategist.
You extract structured professional information from raw career history.

CRITICAL RULES:
- Return ONLY valid JSON. No markdown fences, no explanation, no extra text.
- Follow the EXACT schema provided.
- Summary must be 80-120 words, written in third person professional tone.
- Each role must have MINIMUM 3 achievement bullets using the STAR method.
- Achievements must include quantified results wherever possible.
- Use executive-level language (led, spearheaded, orchestrated, optimized, etc).
- Extract ALL technical and soft skills mentioned or implied.
- Extract all degrees, certifications, and professional development."""

PASS1_USER_PROMPT_TEMPLATE: str = """Analyze the following raw career history and extract structured resume content.

RAW CAREER HISTORY:
\"\"\"
{raw_input}
\"\"\"

Return ONLY a JSON object with this EXACT schema:
{{
  "Summary": "80-120 word professional summary in third person",
  "Experience": [
    {{
      "Role": "Exact job title",
      "Company": "Company name",
      "Duration": "Start - End dates",
      "Achievements": [
        "STAR-method achievement with quantified result 1",
        "STAR-method achievement with quantified result 2",
        "STAR-method achievement with quantified result 3"
      ]
    }}
  ],
  "Skills": ["skill1", "skill2"],
  "Education": ["Degree / Certification details"]
}}

Remember: MINIMUM 3 achievement bullets per role. Quantify everything possible.
Return ONLY the JSON object."""

PASS2_SYSTEM_PROMPT: str = """You are a senior executive resume editor.
You IMPROVE the writing quality of resume JSON content.

CRITICAL RULES:
- Return ONLY valid JSON. No markdown fences, no explanation.
- Do NOT change the schema or rename any keys.
- Do NOT remove any entries.
- ONLY improve the text content within existing values.
- Strengthen action verbs (e.g., managed → orchestrated, helped → facilitated).
- Elevate to executive tone throughout.
- Inject ATS-optimized keywords naturally.
- Ensure all achievements are measurable and specific.
- Fix any grammar or clarity issues.
- Keep Summary between 80-120 words."""

PASS2_USER_PROMPT_TEMPLATE: str = """Improve the writing quality of this resume JSON.
Do NOT change the schema or key names. Only enhance the text values.

Strengthen:
- Action verbs (use executive language)
- ATS keywords
- Measurable achievements
- Grammar and clarity
- Professional tone

INPUT JSON:
{input_json}

Return ONLY the improved JSON object with the identical schema."""

# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title=APP_TITLE,
    page_icon=APP_ICON,
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============================================================
# CUSTOM CSS
# ============================================================

st.markdown("""
<style>
    /* Dark gradient background */
    .stApp {
        background: linear-gradient(160deg, #0f0c29 0%, #1a1a2e 40%, #16213e 100%);
    }

    /* Title block */
    .hero-title {
        text-align: center;
        padding: 2.5rem 0 0.5rem 0;
    }
    .hero-title h1 {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(135deg, #a855f7 0%, #6366f1 50%, #3b82f6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.25rem;
        letter-spacing: -0.5px;
    }
    .hero-title p {
        font-size: 1.15rem;
        color: #94a3b8;
        font-weight: 400;
        margin-top: 0;
    }

    /* Primary action button */
    div.stButton > button {
        background: linear-gradient(135deg, #a855f7 0%, #3b82f6 100%) !important;
        color: white !important;
        border: none !important;
        padding: 0.85rem 2.5rem !important;
        font-size: 1.15rem !important;
        font-weight: 700 !important;
        border-radius: 14px !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        width: 100% !important;
        box-shadow: 0 4px 20px rgba(168, 85, 247, 0.35) !important;
        letter-spacing: 0.3px !important;
    }
    div.stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 8px 30px rgba(168, 85, 247, 0.55) !important;
    }
    div.stButton > button:active {
        transform: translateY(1px) !important;
    }

    /* Download button */
    div.stDownloadButton > button {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
        color: white !important;
        border: none !important;
        padding: 0.75rem 2rem !important;
        font-size: 1.05rem !important;
        font-weight: 600 !important;
        border-radius: 12px !important;
        width: 100% !important;
        box-shadow: 0 4px 15px rgba(16, 185, 129, 0.35) !important;
        transition: all 0.3s ease !important;
    }
    div.stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 25px rgba(16, 185, 129, 0.5) !important;
    }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1e1e2f 0%, #11111d 100%);
        border-right: 1px solid #2d2d44;
    }

    /* Text area styling */
    .stTextArea textarea {
        background-color: #1e1e2f !important;
        color: #e2e8f0 !important;
        border: 1px solid #334155 !important;
        border-radius: 12px !important;
        font-size: 0.95rem !important;
    }
    .stTextArea textarea:focus {
        border-color: #6366f1 !important;
        box-shadow: 0 0 0 2px rgba(99, 102, 241, 0.25) !important;
    }

    /* Expander */
    .streamlit-expanderHeader {
        background-color: #1e1e2f !important;
        border-radius: 10px !important;
        border: 1px solid #334155 !important;
    }

    /* Success / info / warning boxes */
    .stSuccess, .stInfo {
        border-radius: 10px !important;
    }

    /* Divider */
    .section-divider {
        border: none;
        height: 1px;
        background: linear-gradient(90deg, transparent, #6366f1, transparent);
        margin: 1.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# HEADER
# ============================================================

st.markdown(f"""
<div class="hero-title">
    <h1>{APP_ICON} {APP_TITLE}</h1>
    <p>{APP_SUBTITLE}</p>
</div>
<hr class="section-divider">
""", unsafe_allow_html=True)

# ============================================================
# HELPER FUNCTIONS
# ============================================================


def clean_json_text(raw_text: str) -> str:
    """Remove markdown code fences and extraneous whitespace from LLM output."""
    cleaned = raw_text.strip()
    # Strip ```json ... ``` or ``` ... ```
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def parse_json_response(raw_text: str) -> Optional[dict]:
    """Safely parse JSON from an LLM response string."""
    cleaned = clean_json_text(raw_text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None


def validate_resume_data(data: dict) -> dict:
    """Ensure all required keys and types exist in the resume data dict."""
    if not isinstance(data, dict):
        data = {}

    # Summary
    if "Summary" not in data or not isinstance(data["Summary"], str) or not data["Summary"].strip():
        data["Summary"] = "Accomplished professional with a proven track record of delivering results."

    # Experience
    if "Experience" not in data or not isinstance(data["Experience"], list):
        data["Experience"] = []
    validated_experience: list[dict] = []
    for exp in data["Experience"]:
        if not isinstance(exp, dict):
            continue
        validated_exp: dict[str, Any] = {
            "Role": exp.get("Role", "Professional"),
            "Company": exp.get("Company", "Organization"),
            "Duration": exp.get("Duration", ""),
            "Achievements": exp.get("Achievements", []),
        }
        if not isinstance(validated_exp["Achievements"], list):
            validated_exp["Achievements"] = [str(validated_exp["Achievements"])]
        validated_experience.append(validated_exp)
    data["Experience"] = validated_experience

    # Skills
    if "Skills" not in data or not isinstance(data["Skills"], list):
        data["Skills"] = []
    data["Skills"] = [str(s) for s in data["Skills"] if s]

    # Education
    if "Education" not in data or not isinstance(data["Education"], list):
        data["Education"] = []
    data["Education"] = [str(e) for e in data["Education"] if e]

    return data


def call_gemini_with_retry(
    client: genai.Client,
    system_prompt: str,
    user_prompt: str,
    temperature: float,
    status_placeholder: Any,
    step_label: str,
) -> Optional[str]:
    """Call the Gemini model with retry logic (up to MAX_RETRIES attempts)."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            status_placeholder.info(
                f"🔄 {step_label} — Attempt {attempt}/{MAX_RETRIES}..."
            )
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=user_prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    response_mime_type="application/json",
                    temperature=temperature,
                ),
            )

            if response and response.text:
                return response.text

            status_placeholder.warning(
                f"⚠️ {step_label} — Empty response on attempt {attempt}. Retrying..."
            )
        except Exception as exc:
            error_msg = str(exc).lower()

            if "api key" in error_msg or "api_key" in error_msg or "permission" in error_msg:
                st.error("🔑 **Invalid API Key.** Please check your Gemini API key in the sidebar.")
                return None
            elif "quota" in error_msg or "resource_exhausted" in error_msg or "429" in error_msg:
                st.error("📊 **Quota exceeded.** Please wait a moment or check your Gemini plan limits.")
                return None
            elif "timeout" in error_msg or "deadline" in error_msg:
                status_placeholder.warning(
                    f"⏱️ {step_label} — Timeout on attempt {attempt}. Retrying..."
                )
            elif "network" in error_msg or "connection" in error_msg or "unavailable" in error_msg:
                status_placeholder.warning(
                    f"🌐 {step_label} — Network error on attempt {attempt}. Retrying..."
                )
            else:
                status_placeholder.warning(
                    f"⚠️ {step_label} — Error on attempt {attempt}: {exc}. Retrying..."
                )

        if attempt < MAX_RETRIES:
            time.sleep(RETRY_DELAY_SECONDS * attempt)

    st.error(f"❌ {step_label} failed after {MAX_RETRIES} attempts. Please try again.")
    return None


def run_extraction_pass(
    client: genai.Client,
    raw_input: str,
    status_placeholder: Any,
) -> Optional[dict]:
    """PASS 1: Extract structured resume JSON from raw career history."""
    user_prompt = PASS1_USER_PROMPT_TEMPLATE.format(raw_input=raw_input)
    raw_response = call_gemini_with_retry(
        client=client,
        system_prompt=PASS1_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        temperature=0.3,
        status_placeholder=status_placeholder,
        step_label="Pass 1 — Extracting structured content",
    )

    if raw_response is None:
        return None

    parsed = parse_json_response(raw_response)
    if parsed is None:
        st.error("❌ Pass 1 returned invalid JSON. Please try again or rephrase your input.")
        return None

    return validate_resume_data(parsed)


def run_enhancement_pass(
    client: genai.Client,
    extracted_data: dict,
    status_placeholder: Any,
) -> dict:
    """PASS 2: Enhance writing quality of the structured resume JSON."""
    input_json_str = json.dumps(extracted_data, indent=2)
    user_prompt = PASS2_USER_PROMPT_TEMPLATE.format(input_json=input_json_str)
    raw_response = call_gemini_with_retry(
        client=client,
        system_prompt=PASS2_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        temperature=0.2,
        status_placeholder=status_placeholder,
        step_label="Pass 2 — Enhancing executive language",
    )

    if raw_response is None:
        # Fall back to the Pass 1 output
        return extracted_data

    parsed = parse_json_response(raw_response)
    if parsed is None:
        # Fall back to the Pass 1 output
        st.warning("⚠️ Pass 2 returned invalid JSON. Using Pass 1 results.")
        return extracted_data

    return validate_resume_data(parsed)


# ============================================================
# DOCX GENERATION
# ============================================================


def _add_styled_heading(
    doc: Document,
    text: str,
    level: int,
    color: RGBColor,
    size: Optional[int] = None,
    bold: bool = True,
) -> None:
    """Add a heading with custom font color and optional size override."""
    heading = doc.add_heading(level=level)
    run = heading.runs[0] if heading.runs else heading.add_run()
    run.text = text
    run.font.color.rgb = color
    run.bold = bold
    if size:
        run.font.size = Pt(size)


def _add_separator(doc: Document) -> None:
    """Add a thin visual separator line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run("─" * 60)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)


def generate_docx_bytes(resume_data: dict, theme: str) -> bytes:
    """Generate a professionally formatted DOCX resume and return it as bytes.

    Uses python-docx for pure-Python generation with zero system dependencies.
    Theme-dependent heading colors are applied throughout.
    """
    colors = THEME_COLORS.get(theme, THEME_COLORS["Modern-Tech"])
    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # -- Document Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run("PROFESSIONAL RESUME")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = colors["heading"]

    # Theme subtitle
    theme_para = doc.add_paragraph()
    theme_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    theme_para.paragraph_format.space_after = Pt(12)
    theme_run = theme_para.add_run(f"— {theme} Format —")
    theme_run.font.size = Pt(10)
    theme_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    theme_run.italic = True

    _add_separator(doc)

    # -- Professional Summary
    _add_styled_heading(doc, "PROFESSIONAL SUMMARY", level=1, color=colors["heading"], size=14)
    summary_para = doc.add_paragraph(resume_data.get("Summary", ""))
    summary_para.paragraph_format.space_after = Pt(10)
    for run in summary_para.runs:
        run.font.size = Pt(10.5)

    _add_separator(doc)

    # -- Professional Experience
    _add_styled_heading(doc, "PROFESSIONAL EXPERIENCE", level=1, color=colors["heading"], size=14)

    for exp in resume_data.get("Experience", []):
        # Role and Company
        role_para = doc.add_paragraph()
        role_para.paragraph_format.space_before = Pt(8)
        role_para.paragraph_format.space_after = Pt(2)
        role_run = role_para.add_run(f"{exp.get('Role', 'Professional')}")
        role_run.bold = True
        role_run.font.size = Pt(11.5)
        role_run.font.color.rgb = colors["subheading"]

        company_run = role_para.add_run(f"  |  {exp.get('Company', '')}")
        company_run.font.size = Pt(10.5)
        company_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Duration
        duration_text = exp.get("Duration", "")
        if duration_text:
            dur_para = doc.add_paragraph()
            dur_para.paragraph_format.space_before = Pt(0)
            dur_para.paragraph_format.space_after = Pt(4)
            dur_run = dur_para.add_run(duration_text)
            dur_run.font.size = Pt(9.5)
            dur_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            dur_run.italic = True

        # Achievement bullets
        for achievement in exp.get("Achievements", []):
            bullet_para = doc.add_paragraph(style="List Bullet")
            bullet_para.paragraph_format.space_before = Pt(1)
            bullet_para.paragraph_format.space_after = Pt(1)
            bullet_para.paragraph_format.left_indent = Inches(0.35)
            # Clear default run and add styled one
            bullet_para.clear()
            bullet_run = bullet_para.add_run(str(achievement))
            bullet_run.font.size = Pt(10)

    _add_separator(doc)

    # -- Key Skills
    _add_styled_heading(doc, "KEY SKILLS", level=1, color=colors["heading"], size=14)
    skills_list = resume_data.get("Skills", [])
    if skills_list:
        # Render as comma-separated for ATS, or bullets for other themes
        if theme == "ATS-Friendly":
            skills_para = doc.add_paragraph(", ".join(skills_list))
            skills_para.paragraph_format.space_after = Pt(8)
            for run in skills_para.runs:
                run.font.size = Pt(10.5)
        else:
            # Two-column style: pair skills side by side
            for i in range(0, len(skills_list), 2):
                left = skills_list[i]
                right = skills_list[i + 1] if i + 1 < len(skills_list) else ""
                skill_text = f"•  {left}"
                if right:
                    skill_text += f"     •  {right}"
                skill_para = doc.add_paragraph(skill_text)
                skill_para.paragraph_format.space_before = Pt(1)
                skill_para.paragraph_format.space_after = Pt(1)
                for run in skill_para.runs:
                    run.font.size = Pt(10)

    _add_separator(doc)

    # -- Education & Certifications
    _add_styled_heading(doc, "EDUCATION & CERTIFICATIONS", level=1, color=colors["heading"], size=14)
    for edu in resume_data.get("Education", []):
        edu_para = doc.add_paragraph(style="List Bullet")
        edu_para.clear()
        edu_run = edu_para.add_run(str(edu))
        edu_run.font.size = Pt(10.5)

    # -- Write to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# PREVIEW RENDERER
# ============================================================


def render_preview(resume_data: dict) -> None:
    """Render a rich in-app preview of the resume inside an expander."""
    with st.expander("📄 **Preview Generated Resume**", expanded=True):
        # Summary
        st.markdown("### 📋 Professional Summary")
        st.markdown(f"> {resume_data.get('Summary', 'N/A')}")
        st.markdown("---")

        # Experience
        st.markdown("### 💼 Professional Experience")
        for exp in resume_data.get("Experience", []):
            role = exp.get("Role", "Role")
            company = exp.get("Company", "Company")
            duration = exp.get("Duration", "")
            st.markdown(f"**{role}** at *{company}*")
            if duration:
                st.caption(duration)
            for ach in exp.get("Achievements", []):
                st.markdown(f"- {ach}")
            st.markdown("")
        st.markdown("---")

        # Skills
        st.markdown("### 🛠️ Key Skills")
        skills = resume_data.get("Skills", [])
        if skills:
            # Display as tag-like chips
            skill_md = " &nbsp;•&nbsp; ".join(f"**{s}**" for s in skills)
            st.markdown(skill_md)
        st.markdown("---")

        # Education
        st.markdown("### 🎓 Education & Certifications")
        for edu in resume_data.get("Education", []):
            st.markdown(f"- {edu}")


# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.markdown("## ⚙️ Configuration")

api_key: str = st.sidebar.text_input(
    "🔑 Gemini API Key",
    value=os.environ.get("GEMINI_API_KEY", ""),
    type="password",
    help="Enter your Google Gemini API key. Get one at https://aistudio.google.com/apikey",
)

theme: str = st.sidebar.selectbox(
    "🎨 Resume Theme",
    options=THEMES,
    index=0,
    help="Choose the visual style and formatting approach for your resume.",
)

st.sidebar.markdown("---")
st.sidebar.markdown(
    "<small style='color:#64748b;'>Powered by Gemini 2.5 Flash &bull; python-docx</small>",
    unsafe_allow_html=True,
)

# ============================================================
# MAIN INTERFACE
# ============================================================

raw_input: str = st.text_area(
    "📝 Paste your raw career history",
    height=220,
    placeholder=PLACEHOLDER_TEXT,
    help="Include job titles, companies, dates, achievements, skills, education — the more detail, the better the result.",
)

st.markdown("")  # spacer
build_clicked: bool = st.button("🚀 Build Professional Resume")

# ============================================================
# PIPELINE EXECUTION
# ============================================================

if build_clicked:
    # -- Input validation
    if not raw_input.strip():
        st.warning("⚠️ Please paste your career history before building.")
        st.stop()

    if not api_key.strip():
        st.error("🔑 Please enter your Gemini API Key in the sidebar.")
        st.stop()

    # -- Initialize client
    try:
        client = genai.Client(api_key=api_key.strip())
    except Exception as exc:
        st.error(f"❌ Failed to initialize Gemini client: {exc}")
        st.stop()

    # -- Progress tracking
    progress_bar = st.progress(0, text="Initializing pipeline...")
    status = st.empty()

    # ── PASS 1 ──
    progress_bar.progress(10, text="Pass 1: Extracting structured content...")
    with st.spinner("Analyzing career history with Gemini AI..."):
        extracted = run_extraction_pass(client, raw_input, status)

    if extracted is None:
        progress_bar.progress(100, text="Pipeline failed.")
        st.stop()

    progress_bar.progress(45, text="Pass 1 complete ✓")
    status.success("✅ Pass 1 — Structured extraction complete.")

    # ── PASS 2 ──
    progress_bar.progress(50, text="Pass 2: Enhancing executive language...")
    with st.spinner("Elevating resume language to executive quality..."):
        enhanced = run_enhancement_pass(client, extracted, status)

    progress_bar.progress(80, text="Pass 2 complete ✓")
    status.success("✅ Pass 2 — Executive enhancement complete.")

    # ── DOCX GENERATION ──
    progress_bar.progress(85, text="Generating DOCX document...")
    with st.spinner("Building professional document..."):
        try:
            docx_bytes = generate_docx_bytes(enhanced, theme)
        except Exception as exc:
            st.error(f"❌ Document generation failed: {exc}")
            progress_bar.progress(100, text="Pipeline failed.")
            st.stop()

    progress_bar.progress(100, text="🎉 Resume ready!")
    status.empty()

    # ── SUCCESS ──
    st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
    st.success("🎉 **Your professional resume has been built successfully!**")

    # Preview
    render_preview(enhanced)

    # Download
    st.markdown("")
    st.download_button(
        label="📥 Download Resume (.docx)",
        data=docx_bytes,
        file_name="resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
