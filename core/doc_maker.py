import io
import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.prompts import THEME_COLORS

logger = logging.getLogger("resume_architect")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        return ""

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF using PyMuPDF (fitz) if available,
    falling back to pdfminer, then a raw byte-decode attempt."""
    try:
        import fitz  # type: ignore
        with fitz.open(stream=file_bytes, filetype="pdf") as pdf:
            return "\n".join(page.get_text() for page in pdf)
    except ImportError:
        pass

    try:
        from pdfminer.high_level import extract_text as pdfminer_extract  # type: ignore
        return pdfminer_extract(io.BytesIO(file_bytes))
    except ImportError:
        pass

    logger.warning("No PDF parsing library found. Install PyMuPDF or pdfminer.six.")
    return ""

def _rgb(triple: tuple[int, int, int]) -> RGBColor:
    return RGBColor(*triple)

def _add_hyperlink(paragraph: Any, url: str, text: str, color: RGBColor) -> None:
    """Insert a clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), "{:02X}{:02X}{:02X}".format(*color))
    rPr.append(clr)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def _separator(doc: Document, light: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("─" * 65)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC) if light else RGBColor(0xAA, 0xAA, 0xAA)

def _section_heading(doc: Document, text: str, color: RGBColor) -> None:
    h = doc.add_heading(level=1)
    h.clear()
    r = h.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)

def generate_docx_bytes(data: dict, theme: str) -> bytes:
    """Build a themed, professional DOCX resume and return as bytes (BytesIO)."""
    colors = THEME_COLORS.get(theme, THEME_COLORS["Modern-Tech"])
    heading_color = _rgb(colors["heading"])
    sub_color = _rgb(colors["subheading"])
    accent_color = _rgb(colors["accent"])

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin = Inches(0.85)
        sec.right_margin = Inches(0.85)

    # ── NAME / CONTACT HEADER ──────────────────────────────
    name = data.get("Name", "").strip() or "Your Name"
    email = data.get("Email", "").strip()
    phone = data.get("Phone", "").strip()
    linkedin = data.get("LinkedIn", "").strip()
    location = data.get("Location", "").strip()

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(3)
    nr = name_para.add_run(name)
    nr.bold = True
    nr.font.size = Pt(28)
    nr.font.color.rgb = heading_color

    # Contact line
    contact_parts: list[str] = []
    if location:
        contact_parts.append(location)
    if phone:
        contact_parts.append(phone)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(2)

    if contact_parts:
        cr = contact_para.add_run("  |  ".join(contact_parts))
        cr.font.size = Pt(9.5)
        cr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Hyperlinks for email / LinkedIn
    link_para = doc.add_paragraph()
    link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_para.paragraph_format.space_after = Pt(10)

    if email:
        _add_hyperlink(link_para, f"mailto:{email}", email, accent_color)
    if email and linkedin:
        sp = link_para.add_run("   |   ")
        sp.font.size = Pt(9.5)
        sp.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    if linkedin:
        url = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
        _add_hyperlink(link_para, url, linkedin, accent_color)

    _separator(doc)

    # ── PROFESSIONAL SUMMARY ───────────────────────────────
    _section_heading(doc, "Professional Summary", heading_color)
    sp = doc.add_paragraph(data.get("Summary", ""))
    for r in sp.runs:
        r.font.size = Pt(10.5)
    sp.paragraph_format.space_after = Pt(8)

    _separator(doc)

    # ── PROFESSIONAL EXPERIENCE ────────────────────────────
    _section_heading(doc, "Professional Experience", heading_color)

    for exp in data.get("Experience", []):
        role_para = doc.add_paragraph()
        role_para.paragraph_format.space_before = Pt(8)
        role_para.paragraph_format.space_after = Pt(1)
        rr = role_para.add_run(exp.get("Role", "Professional"))
        rr.bold = True
        rr.font.size = Pt(11)
        rr.font.color.rgb = sub_color

        company_run = role_para.add_run(f"  ·  {exp.get('Company', '')}")
        company_run.font.size = Pt(10.5)
        company_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        dur = exp.get("Duration", "")
        if dur:
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after = Pt(3)
            dr = dp.add_run(dur)
            dr.italic = True
            dr.font.size = Pt(9)
            dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        for ach in exp.get("Achievements", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.clear()
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Inches(0.3)
            ar = bp.add_run(str(ach))
            ar.font.size = Pt(10)

    _separator(doc)

    # ── KEY SKILLS  (two-column table for non-ATS themes) ──
    _section_heading(doc, "Key Skills", heading_color)
    skills = data.get("Skills", [])

    if theme == "ATS-Friendly":
        # Plain comma-separated for maximum parser compatibility
        sp2 = doc.add_paragraph(", ".join(skills))
        for r in sp2.runs:
            r.font.size = Pt(10.5)
    else:
        # Two-column table
        if skills:
            half = (len(skills) + 1) // 2
            left_col = skills[:half]
            right_col = skills[half:]
            table = doc.add_table(rows=max(len(left_col), len(right_col)), cols=2)
            table.style = "Table Grid"

            # Remove borders for visual cleanliness
            for row in table.rows:
                for cell in row.cells:
                    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        tc_pr = cell._tc.get_or_add_tcPr()
                        tc_borders = tc_pr.find(qn("w:tcBorders")) or OxmlElement("w:tcBorders")
                        tc_pr.append(tc_borders)
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn("w:val"), "none")
                        border.set(qn("w:sz"), "0")
                        tc_borders.append(border)

            for i, skill in enumerate(left_col):
                cell = table.cell(i, 0)
                cell.text = f"•  {skill}"
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
            for i, skill in enumerate(right_col):
                cell = table.cell(i, 1)
                cell.text = f"•  {skill}"
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

    _separator(doc)

    # ── EDUCATION & CERTIFICATIONS ─────────────────────────
    _section_heading(doc, "Education & Certifications", heading_color)
    for edu in data.get("Education", []):
        ep = doc.add_paragraph(style="List Bullet")
        ep.clear()
        er = ep.add_run(str(edu))
        er.font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def generate_markdown(data: dict) -> str:
    """Generate a clean Markdown version of the resume."""
    lines: list[str] = []
    name = data.get("Name", "").strip()
    if name:
        lines.append(f"# {name}")
    contact = "  |  ".join(filter(None, [
        data.get("Email", ""), data.get("Phone", ""),
        data.get("LinkedIn", ""), data.get("Location", ""),
    ]))
    if contact:
        lines.append(contact)
    lines += ["", "---", "", "## Professional Summary", "", data.get("Summary", ""), "", "---", ""]

    if data.get("Experience"):
        lines.append("## Professional Experience")
        for exp in data["Experience"]:
            lines.append(f"\n### {exp.get('Role')}  ·  {exp.get('Company')}")
            lines.append(f"*{exp.get('Duration', '')}*")
            for ach in exp.get("Achievements", []):
                lines.append(f"- {ach}")

    lines += ["", "---", "", "## Key Skills", ""]
    if data.get("Skills"):
        lines.append(", ".join(data["Skills"]))

    lines += ["", "---", "", "## Education & Certifications", ""]
    for edu in data.get("Education", []):
        lines.append(f"- {edu}")

    return "\n".join(lines)

def generate_ats_text(data: dict) -> str:
    """Generate a plain-text ATS-optimised version of the resume."""
    lines: list[str] = []
    name = data.get("Name", "").strip()
    if name:
        lines += [name.upper(), "=" * len(name)]
    for f in ("Email", "Phone", "LinkedIn", "Location"):
        v = data.get(f, "").strip()
        if v:
            lines.append(f"{f}: {v}")

    lines += ["", "PROFESSIONAL SUMMARY", "-" * 40, data.get("Summary", ""), ""]
    lines += ["PROFESSIONAL EXPERIENCE", "-" * 40]
    for exp in data.get("Experience", []):
        lines.append(f"{exp.get('Role')} | {exp.get('Company')} | {exp.get('Duration', '')}")
        for ach in exp.get("Achievements", []):
            lines.append(f"  * {ach}")
        lines.append("")

    lines += ["KEY SKILLS", "-" * 40]
    lines.append(", ".join(data.get("Skills", [])))
    lines += ["", "EDUCATION & CERTIFICATIONS", "-" * 40]
    for edu in data.get("Education", []):
        lines.append(edu)

    return "\n".join(lines)
